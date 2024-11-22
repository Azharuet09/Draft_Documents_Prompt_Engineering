import os
import openai
from dotenv import load_dotenv
from Constants import GPT_MODEL

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")


def review_transcript(transcript, specific_issues):
    """
    Function to review a DOAH hearing transcript and identify key testimony or evidence
    that was misinterpreted, ignored, or undervalued by the hearing officer.
    
    Args:
        transcript (str): Full transcript of the DOAH hearing.
        specific_issues (str): Specific issues to focus on for review.
    
    Returns:
        str: Insights and findings from GPT-4 regarding the transcript review.
    """
    try:
        prompt = (
            "Please provide me with a detailed response that is approximately 3000 to 4000 words in length."
            "You are an expert legal assistant tasked with reviewing the transcript "
            "of a DOAH hearing. The goal is to identify testimony or evidence that was "
            "misinterpreted, ignored, or undervalued by the hearing officer. Focus on "
            "highlighting inconsistencies between the officer’s findings and the evidence presented. "
            "\n\n"
            "Here is the transcript:\n"
            f"{transcript}\n\n"
            "Specific issues to focus on:\n"
            f"{specific_issues}\n\n"
            "Provide a detailed analysis, including any errors in fact, law, or procedure."
        )
        
        response = openai.ChatCompletion.create(
            model=GPT_MODEL,
            messages=[
                {"role": "system", "content": "You are a legal assistant specializing in appeals for ESE cases."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=12000
        )
        
        answer = response['choices'][0]['message']['content']
        return answer
    except Exception as e:
        return f"Error: {e}"

def explain_legal_framework(case_description):
    """
    Function to dynamically explain the legal framework for appealing a DOAH decision in an ESE case.
    
    Args:
        case_description (str): A brief description of the case, including key facts and issues.

    Returns:
        str: Explanation of federal statutes and key procedural safeguards tailored to the case.
    """
    try:
        prompt = (
            "Please provide me with a detailed response that is approximately 3000 to 4000 words in length."
            "You are an expert legal assistant specializing in appeals for Exceptional Student Education (ESE) cases. "
            "Given the following case description, explain the legal framework applicable to a federal appeal: "
            "\n\n"
            f"Case Description:\n{case_description}\n\n"
            "Discuss the relevant federal statutes, including:\n"
            "1. Individuals with Disabilities Education Act (IDEA).\n"
            "2. Section 504 of the Rehabilitation Act.\n"
            "3. Americans with Disabilities Act (ADA).\n\n"
            "Address procedural safeguards under IDEA, such as the Free Appropriate Public Education (FAPE) standard, "
            "and explain how procedural or substantive violations could affect the student’s rights. "
            "Structure your response to align with the specific issues in the case."
        )
        
        response = openai.ChatCompletion.create(
            model=GPT_MODEL,
            messages=[
                {"role": "system", "content": "You are a legal assistant specializing in appeals for ESE cases."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=12000
        )
        
        answer = response['choices'][0]['message']['content']
        return answer
    except Exception as e:
        return f"Error: {e}"
    

def analyze_case_law(case_description):
    """
    Function to identify and analyze relevant federal case law supporting an appeal
    of a DOAH decision in an ESE case.
    
    Args:
        case_description (str): A brief description of the case, including key facts and issues.
        
    Returns:
        str: Analysis of relevant federal case law, highlighting similarities and legal principles.
    """
    try:
        prompt = (
            "Please provide me with a detailed response that is approximately 3000 to 4000 words in length."
            "You are an expert legal assistant specializing in Exceptional Student Education (ESE) appeals. "
            "Given the following case description, identify and analyze federal case law that supports arguments for overturning the DOAH hearing officer’s decision. "
            "Focus on cases addressing key areas such as Free Appropriate Public Education (FAPE), Least Restrictive Environment (LRE), and procedural compliance. "
            "Compare the factual similarities or legal principles from prior cases to demonstrate why the hearing officer’s decision conflicts with established precedent.\n\n"
            f"Case Description:\n{case_description}\n\n"
            "Provide a structured analysis including:\n"
            "1. Relevant Federal Cases:\n"
            "   - Case Name and Citation\n"
            "   - Brief Summary of Facts\n"
            "   - Key Legal Principles or Holdings\n"
            "2. Comparison to Current Case:\n"
            "   - Similarities in Facts or Legal Issues\n"
            "   - How the Case Law Supports Overturning the Decision\n"
            "3. Conclusion:\n"
            "   - Summary of How the Identified Case Law Undermines the Hearing Officer’s Decision\n"
        )
        
        response = openai.ChatCompletion.create(
            model=GPT_MODEL,
            messages=[
                {"role": "system", "content": "You are a legal assistant specializing in appeals for ESE cases."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=12000
        )
        
        answer = response['choices'][0]['message']['content']
        return answer
    except Exception as e:
        return f"Error: {e}"

def identify_errors_in_law_or_procedure(case_description):
    """
    Function to identify legal and procedural errors in the DOAH hearing officer's decision.
    
    Args:
        case_description (str): A detailed description of the case, including key facts, 
                                issues, and potential legal/procedural errors.
    
    Returns:
        str: Analysis of legal and procedural errors, highlighting their impact on the decision.
    """
    try:
        prompt = (
            "Please provide me with a detailed response that is approximately 3000 to 4000 words in length."
            "You are a legal expert specializing in appeals for Exceptional Student Education (ESE) cases. "
            "Given the following case description, identify specific legal and procedural errors made by the hearing officer. "
            "Focus on:\n"
            "1. Legal Errors:\n"
            "   - Misapplication of legal standards (e.g., Free Appropriate Public Education (FAPE), Least Restrictive Environment (LRE)).\n"
            "   - Reliance on inadmissible evidence or improper interpretation of statutes and case law.\n"
            "   - Failure to properly weigh testimony from expert witnesses.\n"
            "2. Procedural Errors:\n"
            "   - Failure to meet procedural timelines.\n"
            "   - Improper exclusion of evidence.\n"
            "   - Inadequate findings of fact or reasoning in the decision.\n\n"
            "Case Description:\n"
            f"{case_description}\n\n"
            "Provide a structured response as follows:\n"
            "1. Legal Errors:\n"
            "   - Specific Error\n"
            "   - Explanation and Impact on Decision\n"
            "2. Procedural Errors:\n"
            "   - Specific Error\n"
            "   - Explanation and Impact on Decision\n"
            "3. Conclusion:\n"
            "   - Summary of the errors and how they undermine the validity of the decision.\n"
        )
        
        response = openai.ChatCompletion.create(
            model=GPT_MODEL,
            messages=[
                {"role": "system", "content": "You are a legal assistant specializing in identifying errors in law or procedure for ESE appeals."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=12000
        )
        
        answer = response['choices'][0]['message']['content']
        return answer
    except Exception as e:
        return f"Error: {e}"

def apply_standard_of_review(case_issues):
    """
    Function to clarify the federal court's standard of review for administrative decisions 
    and apply it to the specific issues in the case.
    
    Args:
        case_issues (str): A detailed description of the issues being appealed, including 
                           legal and procedural concerns identified earlier.
    
    Returns:
        str: Analysis of the applicable standard of review and its application to the case issues.
    """
    try:
        prompt = (
            "Please provide me with a detailed response that is approximately 3000 to 4000 words in length."
            "You are a legal expert specializing in appeals for Exceptional Student Education (ESE) cases. "
            "Given the issues described, clarify the federal court's standard of review for administrative decisions, "
            "such as the 'modified de novo' standard. Apply this standard to the specific legal and procedural issues "
            "raised in the appeal.\n\n"
            "Focus on:\n"
            "1. Explanation of the 'modified de novo' standard:\n"
            "   - How federal courts review administrative decisions, including deference to factual findings and legal conclusions.\n"
            "   - Differences between this standard and other review standards (e.g., substantial evidence).\n"
            "2. Application to Case Issues:\n"
            "   - Evaluate whether the hearing officer's findings align with or deviate from the proper legal and procedural standards.\n"
            "   - Highlight specific issues where the officer's decision warrants reversal or modification under the standard of review.\n\n"
            "Case Issues:\n"
            f"{case_issues}\n\n"
            "Provide a structured response with:\n"
            "1. Explanation of the Standard of Review.\n"
            "2. Application to Case Issues.\n"
            "3. Conclusion summarizing how the standard supports the appeal.\n"
        )
        
        response = openai.ChatCompletion.create(
            model=GPT_MODEL,
            messages=[
                {"role": "system", "content": "You are a legal assistant specializing in identifying and applying the standard of review for federal appeals."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=12000
        )
        
        answer = response['choices'][0]['message']['content']
        return answer
    except Exception as e:
        return f"Error: {e}"    

def suggest_relief_sought(case_violations):
    """
    Function to suggest potential remedies based on identified legal and procedural violations.
    
    Args:
        case_violations (str): Description of the legal and procedural errors impacting the student's rights.
    
    Returns:
        str: Suggested relief tailored to the violations, including compensatory education, reimbursement, or directives for a new IEP meeting.
    """
    try:
        prompt = (
            "Please provide me with a detailed response that is approximately 3000 to 4000 words in length."
            "You are a legal expert specializing in appeals for Exceptional Student Education (ESE) cases. "
            "Given the violations described, suggest potential remedies that address the legal and procedural errors. "
            "The remedies must align with IDEA, Section 504, and related federal statutes.\n\n"
            "Focus on:\n"
            "1. Explanation of each suggested remedy and its purpose.\n"
            "2. How each remedy directly addresses the violations in the case.\n\n"
            "Potential remedies may include:\n"
            "- Compensatory education to make up for lost educational opportunities.\n"
            "- Reimbursement for private school tuition or services.\n"
            "- An order for the school district to convene a new IEP team meeting with specific directives.\n\n"
            "Case Violations:\n"
            f"{case_violations}\n\n"
            "Provide a structured response with:\n"
            "1. Overview of the violations.\n"
            "2. Suggested Remedies.\n"
            "3. Justification for each remedy.\n"
        )
        
        response = openai.ChatCompletion.create(
            model=GPT_MODEL,
            messages=[
                {"role": "system", "content": "You are a legal assistant specializing in suggesting appropriate remedies in federal appeals for ESE cases."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=12000
        )
        
        answer = response['choices'][0]['message']['content']
        return answer
    except Exception as e:
        return f"Error: {e}"

def draft_persuasive_argument(case_details, counterarguments):
    """
    Function to draft a persuasive legal argument for a federal appeal in an ESE case.
    
    Args:
        case_details (str): Key facts and legal issues in the case.
        counterarguments (str): Anticipated counterarguments from the school district.
    
    Returns:
        str: A well-written argument emphasizing the equities and addressing counterarguments.
    """
    try:
        prompt = (
            "Please provide me with a detailed response that is approximately 3000 to 4000 words in length."
            "You are a legal writing expert specializing in federal appeals for Exceptional Student Education (ESE) cases. "
            "Draft a persuasive argument for an appeal using the following guidelines:\n\n"
            "1. Use a clear, concise, and structured writing style.\n"
            "2. Anticipate and refute counterarguments raised by the school district.\n"
            "3. Emphasize the equities involved, particularly the impact of the decision on the student’s education.\n"
            "4. Focus on the legal and factual grounds for overturning the hearing officer’s decision.\n\n"
            "Key Case Details:\n"
            f"{case_details}\n\n"
            "Anticipated Counterarguments:\n"
            f"{counterarguments}\n\n"
            "Draft a structured response with:\n"
            "1. Introduction summarizing the appeal and its purpose.\n"
            "2. Detailed argument refuting the school district’s counterarguments.\n"
            "3. Emphasis on the impact of the decision on the student’s education.\n"
            "4. Conclusion reinforcing why the decision should be overturned."
        )
        
        response = openai.ChatCompletion.create(
            model=GPT_MODEL,
            messages=[
                {"role": "system", "content": "You are a legal expert assisting in drafting persuasive appellate arguments for ESE cases."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=12000
        )
        
        answer = response['choices'][0]['message']['content']
        return answer
    except Exception as e:
        return f"Error: {e}"

from docx import Document
import math

def create_pages_from_result(final_result, words_per_page=300, doc_filename="case_analysis.docx"):
    """
    Create a multi-page document (50-70 pages) from the final result string.
    
    :param final_result: The complete text to be added to the document.
    :param words_per_page: The target word count per page for splitting.
    :param doc_filename: The name of the output .docx file.
    """
    # Create a new document
    doc = Document()
    
    # Split final result into words
    words = final_result.split()
    total_words = len(words)
    
    # Calculate number of pages needed
    pages_needed = math.ceil(total_words / words_per_page)
    
    # Add content to the document, splitting based on words_per_page
    current_page = 1
    start_index = 0
    while start_index < total_words:
        # Define the end index for this page
        end_index = min(start_index + words_per_page, total_words)
        
        # Create a new page
        # doc.add_paragraph("Page %d" % current_page, style="Heading 1")
        
        # Add the portion of text for this page
        page_text = " ".join(words[start_index:end_index])
        doc.add_paragraph(page_text)
        
        # Prepare for the next page
        start_index = end_index
        current_page += 1
        
        # Add a page break after each page except the last one
        if start_index < total_words:
            doc.add_paragraph("\n")
            doc.add_paragraph("\n")
    
    # Save the document
    doc.save(doc_filename)
    print(f"Document saved as {doc_filename}. It contains {pages_needed} pages.")

if __name__ == "__main__":
    transcript = "Full DOAH hearing transcript goes here."
    specific_issues = "Highlight testimony related to student assessments and procedural compliance."
    
    result = review_transcript(transcript, specific_issues)
    case_description = (
        "The student is a 10-year-old with autism who was denied necessary speech therapy services. "
        "The hearing officer ruled that the school district's IEP was sufficient, but the family contends "
        "that the IEP fails to meet the Free Appropriate Public Education (FAPE) standard. Procedural violations, "
        "such as failure to consider independent evaluations, were also noted."
    )
    
    result1 = explain_legal_framework(case_description)
    result2 = analyze_case_law(case_description)
    case_description1 = (
        "The hearing officer denied speech therapy services for a 10-year-old student with autism, "
        "ruling that the existing IEP met FAPE standards. However, the officer relied on testimony "
        "from a school district employee who lacked expertise in speech therapy, ignored independent evaluations, "
        "and dismissed the parents' evidence as inadmissible. Procedurally, the school district failed to meet IDEA "
        "timelines for responding to the parents' requests for evaluations."
    )
    
    result3 = identify_errors_in_law_or_procedure(case_description1)
    case_issues = (
        "The hearing officer ruled that the student's IEP met the FAPE standard, despite substantial evidence from "
        "independent evaluations showing lack of meaningful progress in speech therapy. The officer also excluded "
        "critical expert testimony presented by the parents and failed to address procedural violations, such as "
        "the district's delay in responding to evaluation requests."
    )
    
    result4 = apply_standard_of_review(case_issues)
    case_violations = (
        "The school district failed to provide FAPE by not addressing the student's speech therapy needs, despite "
        "independent evaluations showing a lack of progress. The hearing officer also excluded critical expert testimony "
        "and overlooked procedural violations, including delays in responding to evaluation requests. As a result, the "
        "student has fallen behind significantly in communication skills."
    )
    
    result5 = suggest_relief_sought(case_violations)

    case_details = (
        "The hearing officer incorrectly concluded that the district provided FAPE despite clear evidence of "
        "academic regression, inadequate IEP implementation, and exclusion of critical expert testimony. "
        "The decision failed to address procedural violations that denied the student’s rights under IDEA."
    )
    
    counterarguments = (
        "The school district may argue that the IEP was reasonably calculated to provide educational benefits "
        "and that any procedural violations were harmless. They may also assert that the parent's disagreement "
        "does not invalidate the IEP."
    )
    
    result6 = draft_persuasive_argument(case_details, counterarguments)

    # final_result = result + result1 + result2 + result3 + result4 + result5 + result6
    # Assuming final_result is already generated from the previous results
    final_result = result + result1 + result2 + result3 + result4 + result5 + result6

    # Call the function to create the pages
    create_pages_from_result(final_result, words_per_page=300, doc_filename="case_analysis.docx")